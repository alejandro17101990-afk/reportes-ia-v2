import streamlit as st
from docx import Document
import speech_recognition as sr
from openai import OpenAI
import io

# 1. CONFIGURACIÓN DEL WORKSPACE
st.set_page_config(page_title="Beam AI Radiology", layout="wide", initial_sidebar_state="expanded")

# 2. MOTOR UI/UX: ESTÉTICA PREMIUM
st.markdown("""
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600&display=swap');
    .stApp { background-color: #0b0b0f; color: #e2e8f0; font-family: 'Inter', sans-serif; }
    [data-testid="stSidebar"] { background-color: #111116; border-right: 1px solid #1f1f2e; }
    .stButton>button { 
        background: linear-gradient(135deg, #7c3aed 0%, #6d28d9 100%) !important;
        color: white !important; border-radius: 8px !important; border: 1px solid #8b5cf6 !important;
        padding: 0.6rem 1rem !important; font-weight: 500 !important; width: 100%;
    }
    div[data-testid="stVerticalBlock"] > div[style*="flex-direction: column;"] { 
        background-color: #12121a; border: 1px solid #1f1f2e; border-radius: 12px; padding: 24px; 
    }
    .stTextInput input, .stSelectbox div[data-baseweb="select"], .stTextArea textarea { 
        background-color: #16161d !important; color: #f8fafc !important; 
        border: 1px solid #2a2a35 !important; border-radius: 8px !important; 
    }
    h1, h2, h3, h4 { color: #ffffff !important; font-weight: 600 !important; }
    hr { border-color: #1f1f2e !important; }
    </style>
    """, unsafe_allow_html=True)

# 3. FUNCIONES DEL NÚCLEO
def leer_word(file):
    doc = Document(file)
    return '\n'.join([para.text for para in doc.paragraphs])

def generar_docx(texto_limpio):
    doc = Document()
    doc.add_paragraph(texto_limpio)
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

def transcribir_audio(audio_file):
    r = sr.Recognizer()
    with sr.AudioFile(audio_file) as source:
        audio_data = r.record(source)
        try:
            return r.recognize_google(audio_data, language="es-MX")
        except sr.UnknownValueError:
            return "[No se detectó voz clara]"
        except Exception as e:
            return f"[Error de audio: {e}]"

if 'reporte_generado' not in st.session_state:
    st.session_state.reporte_generado = ""

# 4. BARRA LATERAL
with st.sidebar:
    st.markdown("### ⚕️ Beam AI")
    try:
        api_key = st.secrets["deepseek_key"]
        st.success("Motor DeepSeek: Conectado 🟢")
    except Exception:
        api_key = st.text_input("DeepSeek Key", type="password")
    
    st.divider()
    archivo_plantilla = st.file_uploader("Cargar Plantilla Base (.docx)", type=["docx"])
    plantilla_contenido = leer_word(archivo_plantilla) if archivo_plantilla else ""

# 5. ESPACIO DE TRABAJO PRINCIPAL
if api_key:
    # Aumentamos el tiempo de espera a 45 segundos por si el servidor está lento
    client = OpenAI(api_key=api_key, base_url="https://api.deepseek.com", timeout=45.0)

    st.markdown("## ⚡ Nueva Interpretación")
    col_input, col_output = st.columns([1, 1.2], gap="large")

    with col_input:
        st.markdown("#### 🎙️ Entrada de Datos")
        mod_col, sub_col = st.columns(2)
        with mod_col:
            modalidad = st.selectbox("Estudio", ["TC Lumbar", "RM Rodilla", "TC Cráneo", "Radiografía"])
        with sub_col:
            estilo = st.selectbox("Estilo", ["Conciso", "Académico", "Institucional"])
        
        audio_file = st.audio_input("Iniciar dictado (Opcional)")
        notas_texto = st.text_area("Notas manuales:", height=150, placeholder="Ej. L4-L5 con abombamiento discal difuso...")

        if st.button("✨ Generar Informe Clínico"):
            texto_dictado = ""
            if audio_file:
                with st.spinner("Procesando voz..."):
                    texto_dictado = transcribir_audio(audio_file)
                    st.info(f"Voz: {texto_dictado}")

            if not notas_texto and not texto_dictado:
                st.warning("⚠️ Ingresa hallazgos por texto o voz antes de generar.")
            else:
                prompt_sistema = f"""
                Eres un Médico Radiólogo experto. Genera un informe estructurado de {modalidad} en estilo {estilo}.
                Estructura: Técnica, Hallazgos, Impresión Diagnóstica.
                No uses código HTML ni asteriscos excesivos, redacta texto limpio y formal.
                PLANTILLA BASE: {plantilla_contenido if plantilla_contenido else "Ninguna"}
                """
                
                prompt_usuario = f"Genera el reporte con estos hallazgos: Notas: {notas_texto} | Dictado: {texto_dictado}"

                with st.spinner("🧠 DeepSeek analizando hallazgos..."):
                    try:
                        response = client.chat.completions.create(
                            model="deepseek-chat",
                            messages=[
                                {"role": "system", "content": prompt_sistema},
                                {"role": "user", "content": prompt_usuario}
                            ],
                            temperature=0.2
                        )
                        st.session_state.reporte_generado = response.choices[0].message.content
                    except Exception as e:
                        # Si falla, ¡esto nos dirá exactamente por qué!
                        st.error(f"❌ Error de conexión con la IA: {str(e)}")

    with col_output:
        st.markdown("#### 📄 Reporte Final")
        
        texto_final = st.text_area(
            "Revisa y edita:", 
            value=st.session_state.reporte_generado, 
            height=450
        )
        
        if st.session_state.reporte_generado:
            st.download_button(
                label="📥 Descargar a Word", 
                data=generar_docx(texto_final), 
                file_name="Reporte_Radiologico.docx", 
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
else:
    st.info("Configura tu credencial para inicializar el Workspace.")